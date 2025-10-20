from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from typing import Dict
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json, re

app = FastAPI(title="WorkDocStyler API")

# --- Arkance Style Map: default embedded (you can override via style_map_json) ---
DEFAULT_STYLE_MAP = {
  "Heading 1": {
    "font_name": "Century Gothic", "font_size_pt": 14, "bold": True, "italic": False,
    "color": "#0052A3", "alignment": "Left", "line_spacing": 1.15,
    "spacing_before_pt": 12, "spacing_after_pt": 6, "keep_with_next": True,
    "keep_lines_together": True, "indent_left_cm": 0, "indent_hanging_cm": 1,
    "numbering_level": 1, "numbering_pattern": "%1", "based_on": "Normal",
    "following_style": "Normal"
  },
  "Heading 2": {
    "font_name": "Century Gothic", "font_size_pt": 12, "bold": True, "italic": False,
    "color": "#0052A3", "alignment": "Left", "line_spacing": 1.15,
    "spacing_before_pt": 12, "spacing_after_pt": 6, "keep_with_next": True,
    "keep_lines_together": True, "indent_left_cm": 0, "indent_hanging_cm": 1.02,
    "numbering_level": 2, "numbering_pattern": "%1.%2", "based_on": "Normal",
    "following_style": "Normal"
  },
  "Heading 3": {
    "font_name": "Century Gothic", "font_size_pt": 11, "bold": True, "italic": False,
    "color": "#0052A3", "alignment": "Left", "line_spacing": 1.15,
    "spacing_before_pt": 8, "spacing_after_pt": 6, "keep_with_next": True,
    "keep_lines_together": True, "indent_left_cm": 0, "indent_hanging_cm": 1.27,
    "numbering_level": 3, "numbering_pattern": "%1.%2.%3", "based_on": "Normal",
    "following_style": "Normal"
  },
  "Heading 4": {
    "font_name": "Century Gothic", "font_size_pt": 9, "bold": False, "italic": False,
    "color": "RGB(1,95,95)", "alignment": "Left", "line_spacing": 1.5,
    "spacing_before_pt": 8, "spacing_after_pt": 6, "keep_with_next": True,
    "keep_lines_together": True, "indent_left_cm": 0, "indent_hanging_cm": 1.52,
    "numbering_level": 4, "numbering_pattern": "%1.%2.%3.%4", "based_on": "Normal",
    "following_style": "Normal"
  },
  "Normal": {
    "font_name": "Century Gothic", "font_size_pt": 10, "bold": False, "italic": False,
    "color": "Text 1", "alignment": "Left", "line_spacing": 1.0,
    "spacing_before_pt": 0, "spacing_after_pt": 6, "widow_orphan_control": True,
    "based_on": None, "following_style": "Normal"
  },
  "List Paragraph Bullet Points": {
    "font_name": "Century Gothic", "font_size_pt": 10.5, "bold": False, "italic": False,
    "color": "Text 1", "alignment": "Left", "line_spacing": 1.0,
    "spacing_before_pt": 0, "spacing_after_pt": 0, "indent_left_cm": 1.27,
    "spacing_same_paragraphs": False, "based_on": "Normal",
    "following_style": "List Paragraph Bullet Points"
  },
  "Normal Bullet": {
    "font_name": "Century Gothic", "font_size_pt": 10, "bold": False, "italic": False,
    "color": "Text 1", "alignment": "Left", "line_spacing": 1.08,
    "spacing_before_pt": 8, "spacing_after_pt": 8, "indent_hanging_cm": 0.63,
    "indent_left_cm": 1.27, "bullet_level": 1, "bullet_alignment_cm": 0.63,
    "based_on": "List Paragraph Bullet Point", "following_style": "Normal Bullet"
  }
}

def _rgb(color_str: str):
    if isinstance(color_str, str) and color_str.startswith("#") and len(color_str)==7:
        return RGBColor(int(color_str[1:3],16), int(color_str[3:5],16), int(color_str[5:7],16))
    if isinstance(color_str,str) and color_str.upper().startswith("RGB("):
        parts = color_str[4:-1].split(",")
        try:
            r,g,b = [int(x.strip()) for x in parts]
            return RGBColor(r,g,b)
        except:
            pass
    return RGBColor(0,0,0)  # fallback for "Text 1/2" etc.

def _apply_paragraph_style(p, spec: Dict):
    align_map = {"Left": WD_ALIGN_PARAGRAPH.LEFT, "Center": WD_ALIGN_PARAGRAPH.CENTER,
                 "Right": WD_ALIGN_PARAGRAPH.RIGHT, "Justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    if spec.get("alignment") in align_map:
        p.alignment = align_map[spec["alignment"]]
    if "line_spacing" in spec: p.paragraph_format.line_spacing = spec["line_spacing"]
    if "spacing_before_pt" in spec: p.paragraph_format.space_before = Pt(spec["spacing_before_pt"])
    if "spacing_after_pt" in spec: p.paragraph_format.space_after = Pt(spec["spacing_after_pt"])
    if "indent_left_cm" in spec: p.paragraph_format.left_indent = Cm(spec["indent_left_cm"])
    if "indent_hanging_cm" in spec: p.paragraph_format.first_line_indent = -Cm(spec["indent_hanging_cm"])
    if spec.get("keep_with_next"): p.paragraph_format.keep_with_next = True
    if spec.get("keep_lines_together"): p.paragraph_format.keep_together = True

    runs = p.runs or [p.add_run("")]
    for r in runs:
        if "font_name" in spec: r.font.name = spec["font_name"]
        if "font_size_pt" in spec: r.font.size = Pt(spec["font_size_pt"])
        if "bold" in spec: r.font.bold = spec["bold"]
        if "italic" in spec: r.font.italic = spec["italic"]
        if "color" in spec: r.font.color.rgb = _rgb(spec["color"])

def _add_styled_paragraph(doc: Document, text: str, style_name: str, rules: Dict, counters: Dict):
    p = doc.add_paragraph(text)
    _apply_paragraph_style(p, rules.get(style_name, {}))
    counters[style_name] = counters.get(style_name, 0) + 1

def _detect_and_style_line(doc, line: str, rules: Dict, counters: Dict):
    # Normalize line and strip any UTF-8 BOM if present (common on first line)
    text = (line or "").rstrip("\r\n")
    if text.startswith("\ufeff"):
        text = text[1:]

    # Headings by markers
    heading_markers = [
        ("Heading 1", ["H1:", "# "]),
        ("Heading 2", ["H2:", "## "]),
        ("Heading 3", ["H3:", "### "]),
        ("Heading 4", ["H4:", "#### "]),
    ]
    for style_name, prefixes in heading_markers:
        for pref in prefixes:
            if text.startswith(pref):
                clean = text[len(pref):].strip()
                _add_styled_paragraph(doc, clean, style_name, rules, counters)
                return

    # Numbered lists (simple detection: "1. ", "1) ", "a. ", "A) ")
    numbered_re = r"^(\d+[\.\)]\s|[A-Za-z][\.\)]\s)"
    if re.match(numbered_re, text):
        clean = re.sub(numbered_re, "", text).strip()
        _add_styled_paragraph(doc, clean, "Normal", rules, counters)
        return

    # Bullets (- / * / • )
    if any(text.startswith(m) for m in ["- ", "* ", "• "]):
        clean = text[2:].strip() if len(text) >= 2 else ""
        style_choice = "Normal Bullet" if "Normal Bullet" in rules else "List Paragraph Bullet Points"
        _add_styled_paragraph(doc, clean, style_choice, rules, counters)
        return

    # Default = Normal
    _add_styled_paragraph(doc, text, "Normal", rules, counters)

@app.post("/format")
async def format_doc(
    draft: UploadFile = File(..., description=".txt or .docx"),
    style_map_json: str = Form(None, description="Optional: override style map")
):
    # Decide which rules to use
    rules = DEFAULT_STYLE_MAP
    if style_map_json:
        try:
            rules = json.loads(style_map_json)
        except:
            raise HTTPException(400, "Invalid style_map_json JSON")

    # Read input file
    data = await draft.read()
    out_doc = Document()

    if draft.filename.lower().endswith(".txt"):
        lines = data.decode("utf-8", errors="ignore").splitlines()
    elif draft.filename.lower().endswith(".docx"):
        src = Document(BytesIO(data))
        lines = [p.text for p in src.paragraphs]
    else:
        raise HTTPException(400, "Unsupported file (use .txt or .docx)")

    # Build output
    counters: Dict[str,int] = {}
    for line in lines:
        _detect_and_style_line(out_doc, line, rules, counters)

    # Stream .docx back
    buf = BytesIO()
    out_doc.save(buf)
    buf.seek(0)
    headers = {"X-Delta-Report": json.dumps(counters)}
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )
